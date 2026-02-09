"""
Document RAG (Retrieval Augmented Generation) for Enigma AI Engine

Process and retrieve information from documents.

Features:
- Multi-format document loading (PDF, DOCX, TXT, Markdown)
- Chunking with overlap
- Vector embedding and retrieval
- Hybrid search (keyword + semantic)
- Source citation

Usage:
    from enigma_engine.memory.document_rag import DocumentRAG, Document
    
    # Create RAG system
    rag = DocumentRAG()
    
    # Add documents
    rag.add_file("report.pdf")
    rag.add_file("notes.txt")
    
    # Query
    results = rag.query("What are the key findings?")
    
    # Generate answer with context
    answer = rag.answer("Summarize the main points", model=my_model)
"""

import hashlib
import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterator, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A document in the RAG system."""
    id: str
    content: str
    source: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    # Chunking info
    chunk_index: int = 0
    total_chunks: int = 1
    start_char: int = 0
    end_char: int = 0
    
    # Embedding (computed later)
    embedding: Optional[List[float]] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "id": self.id,
            "content": self.content,
            "source": self.source,
            "metadata": self.metadata,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks
        }


@dataclass
class SearchResult:
    """A search result."""
    document: Document
    score: float
    match_type: str = "semantic"  # semantic, keyword, hybrid
    highlights: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "score": self.score,
            "match_type": self.match_type,
            "source": self.document.source,
            "content": self.document.content[:500],
            "highlights": self.highlights
        }


@dataclass
class RAGConfig:
    """RAG system configuration."""
    # Chunking
    chunk_size: int = 512
    chunk_overlap: int = 50
    
    # Search
    top_k: int = 5
    min_score: float = 0.5
    
    # Hybrid search weights
    semantic_weight: float = 0.7
    keyword_weight: float = 0.3
    
    # Embedding
    embedding_model: str = "default"
    
    # Context building
    max_context_tokens: int = 2000
    include_source: bool = True


class DocumentLoader:
    """Loads documents from various formats."""
    
    @staticmethod
    def load(path: str) -> str:
        """
        Load document from file.
        
        Args:
            path: File path
            
        Returns:
            Document text content
        """
        path = Path(path)
        suffix = path.suffix.lower()
        
        if suffix == ".txt" or suffix == ".md":
            return DocumentLoader._load_text(path)
        elif suffix == ".pdf":
            return DocumentLoader._load_pdf(path)
        elif suffix == ".docx":
            return DocumentLoader._load_docx(path)
        elif suffix == ".html":
            return DocumentLoader._load_html(path)
        elif suffix == ".json":
            return DocumentLoader._load_json(path)
        else:
            # Try as text
            return DocumentLoader._load_text(path)
    
    @staticmethod
    def _load_text(path: Path) -> str:
        """Load text file."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    @staticmethod
    def _load_pdf(path: Path) -> str:
        """Load PDF file."""
        try:
            import PyPDF2
            
            text = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            
            return "\n\n".join(text)
            
        except ImportError:
            logger.warning("PyPDF2 not available, trying pdfplumber")
            
            try:
                import pdfplumber
                
                text = []
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text.append(page.extract_text() or "")
                
                return "\n\n".join(text)
                
            except ImportError:
                logger.error("No PDF library available (install PyPDF2 or pdfplumber)")
                return ""
    
    @staticmethod
    def _load_docx(path: Path) -> str:
        """Load DOCX file."""
        try:
            import docx
            
            doc = docx.Document(str(path))
            return "\n\n".join(para.text for para in doc.paragraphs)
            
        except ImportError:
            logger.error("python-docx not available")
            return ""
    
    @staticmethod
    def _load_html(path: Path) -> str:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(path, encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            
            # Remove scripts and styles
            for script in soup(["script", "style"]):
                script.extract()
            
            return soup.get_text(separator="\n")
            
        except ImportError:
            # Fall back to simple regex
            with open(path, encoding="utf-8") as f:
                html = f.read()
            return re.sub(r"<[^>]+>", "", html)
    
    @staticmethod
    def _load_json(path: Path) -> str:
        """Load JSON file as text."""
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)


class TextChunker:
    """Chunks text into smaller pieces."""
    
    def __init__(
        self,
        chunk_size: int = 512,
        overlap: int = 50,
        separator: str = "\n"
    ):
        """
        Initialize chunker.
        
        Args:
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks
            separator: Preferred split character
        """
        self._chunk_size = chunk_size
        self._overlap = overlap
        self._separator = separator
    
    def chunk(self, text: str) -> List[Tuple[str, int, int]]:
        """
        Chunk text.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of (chunk_text, start_char, end_char)
        """
        if len(text) <= self._chunk_size:
            return [(text, 0, len(text))]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find end of chunk
            end = start + self._chunk_size
            
            if end >= len(text):
                # Last chunk
                chunks.append((text[start:], start, len(text)))
                break
            
            # Try to split at separator
            split_pos = text.rfind(self._separator, start, end)
            
            if split_pos <= start:
                # No separator found, split at chunk_size
                split_pos = end
            else:
                split_pos += len(self._separator)
            
            chunks.append((text[start:split_pos], start, split_pos))
            
            # Move start with overlap
            start = split_pos - self._overlap
            if start < 0:
                start = 0
        
        return chunks


class SimpleEmbedder:
    """Simple embedding using TF-IDF or sentence transformers."""
    
    def __init__(self, model_name: str = "default"):
        """Initialize embedder."""
        self._model_name = model_name
        self._model = None
        
        # Try to load sentence transformers
        try:
            from sentence_transformers import SentenceTransformer
            self._model = SentenceTransformer("all-MiniLM-L6-v2")
            logger.info("Using sentence-transformers for embeddings")
        except ImportError:
            logger.warning("sentence-transformers not available, using TF-IDF")
    
    def embed(self, texts: List[str]) -> List[List[float]]:
        """
        Embed texts.
        
        Args:
            texts: List of texts
            
        Returns:
            List of embeddings
        """
        if self._model:
            embeddings = self._model.encode(texts)
            return embeddings.tolist()
        else:
            # Simple TF-IDF-like embedding
            return [self._tfidf_embed(t) for t in texts]
    
    def _tfidf_embed(self, text: str, dim: int = 384) -> List[float]:
        """Simple hash-based embedding."""
        words = text.lower().split()
        embedding = [0.0] * dim
        
        for word in words:
            h = int(hashlib.md5(word.encode()).hexdigest(), 16)
            idx = h % dim
            embedding[idx] += 1.0
        
        # Normalize
        norm = sum(x**2 for x in embedding) ** 0.5
        if norm > 0:
            embedding = [x / norm for x in embedding]
        
        return embedding


class DocumentRAG:
    """
    Document-based RAG system.
    """
    
    def __init__(
        self,
        config: Optional[RAGConfig] = None,
        persist_dir: Optional[str] = None
    ):
        """
        Initialize RAG system.
        
        Args:
            config: RAG configuration
            persist_dir: Directory to persist data
        """
        self._config = config or RAGConfig()
        self._persist_dir = Path(persist_dir) if persist_dir else None
        
        self._documents: Dict[str, Document] = {}
        self._loader = DocumentLoader()
        self._chunker = TextChunker(
            chunk_size=self._config.chunk_size,
            overlap=self._config.chunk_overlap
        )
        self._embedder = SimpleEmbedder(self._config.embedding_model)
        
        # Keyword index
        self._keyword_index: Dict[str, List[str]] = {}  # word -> [doc_ids]
        
        # Load persisted data
        if self._persist_dir and self._persist_dir.exists():
            self._load()
    
    def add_file(
        self,
        path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Add a document file.
        
        Args:
            path: File path
            metadata: Optional metadata
            
        Returns:
            List of created document IDs
        """
        path = Path(path)
        
        # Load content
        content = self._loader.load(str(path))
        
        if not content:
            logger.warning(f"Empty content from: {path}")
            return []
        
        # Chunk content
        chunks = self._chunker.chunk(content)
        
        # Create documents
        doc_ids = []
        base_id = hashlib.md5(str(path).encode()).hexdigest()[:8]
        
        for i, (chunk_text, start, end) in enumerate(chunks):
            doc_id = f"{base_id}_{i}"
            
            doc = Document(
                id=doc_id,
                content=chunk_text,
                source=str(path),
                metadata={
                    "filename": path.name,
                    "file_type": path.suffix,
                    **(metadata or {})
                },
                chunk_index=i,
                total_chunks=len(chunks),
                start_char=start,
                end_char=end
            )
            
            # Compute embedding
            embeddings = self._embedder.embed([chunk_text])
            doc.embedding = embeddings[0]
            
            self._documents[doc_id] = doc
            self._index_keywords(doc)
            doc_ids.append(doc_id)
        
        logger.info(f"Added {len(doc_ids)} chunks from: {path}")
        
        # Persist if configured
        if self._persist_dir:
            self._save()
        
        return doc_ids
    
    def add_text(
        self,
        text: str,
        source: str = "direct",
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """Add text directly."""
        chunks = self._chunker.chunk(text)
        doc_ids = []
        base_id = hashlib.md5(text[:100].encode()).hexdigest()[:8]
        
        for i, (chunk_text, start, end) in enumerate(chunks):
            doc_id = f"{base_id}_{i}"
            
            doc = Document(
                id=doc_id,
                content=chunk_text,
                source=source,
                metadata=metadata or {},
                chunk_index=i,
                total_chunks=len(chunks)
            )
            
            embeddings = self._embedder.embed([chunk_text])
            doc.embedding = embeddings[0]
            
            self._documents[doc_id] = doc
            self._index_keywords(doc)
            doc_ids.append(doc_id)
        
        return doc_ids
    
    def _index_keywords(self, doc: Document):
        """Index document keywords."""
        words = set(doc.content.lower().split())
        for word in words:
            if len(word) > 2:
                if word not in self._keyword_index:
                    self._keyword_index[word] = []
                self._keyword_index[word].append(doc.id)
    
    def query(
        self,
        query: str,
        top_k: Optional[int] = None,
        search_type: str = "hybrid"
    ) -> List[SearchResult]:
        """
        Query the document collection.
        
        Args:
            query: Search query
            top_k: Number of results
            search_type: "semantic", "keyword", or "hybrid"
            
        Returns:
            List of search results
        """
        top_k = top_k or self._config.top_k
        
        if search_type == "semantic":
            results = self._semantic_search(query, top_k)
        elif search_type == "keyword":
            results = self._keyword_search(query, top_k)
        else:  # hybrid
            results = self._hybrid_search(query, top_k)
        
        return results
    
    def _semantic_search(self, query: str, top_k: int) -> List[SearchResult]:
        """Semantic search using embeddings."""
        # Embed query
        query_emb = self._embedder.embed([query])[0]
        
        # Compute similarities
        scores = []
        for doc_id, doc in self._documents.items():
            if doc.embedding:
                score = self._cosine_similarity(query_emb, doc.embedding)
                scores.append((doc, score))
        
        # Sort and return top k
        scores.sort(key=lambda x: x[1], reverse=True)
        
        results = []
        for doc, score in scores[:top_k]:
            if score >= self._config.min_score:
                results.append(SearchResult(
                    document=doc,
                    score=score,
                    match_type="semantic"
                ))
        
        return results
    
    def _keyword_search(self, query: str, top_k: int) -> List[SearchResult]:
        """Keyword-based search."""
        query_words = set(query.lower().split())
        
        # Count matches per document
        doc_scores: Dict[str, int] = {}
        for word in query_words:
            if word in self._keyword_index:
                for doc_id in self._keyword_index[word]:
                    doc_scores[doc_id] = doc_scores.get(doc_id, 0) + 1
        
        # Sort by score
        sorted_docs = sorted(
            doc_scores.items(),
            key=lambda x: x[1],
            reverse=True
        )
        
        results = []
        for doc_id, score in sorted_docs[:top_k]:
            if doc_id in self._documents:
                # Normalize score
                norm_score = score / len(query_words) if query_words else 0
                
                # Find highlights
                highlights = []
                for word in query_words:
                    if word.lower() in self._documents[doc_id].content.lower():
                        highlights.append(word)
                
                results.append(SearchResult(
                    document=self._documents[doc_id],
                    score=norm_score,
                    match_type="keyword",
                    highlights=highlights
                ))
        
        return results
    
    def _hybrid_search(self, query: str, top_k: int) -> List[SearchResult]:
        """Combined semantic + keyword search."""
        semantic_results = self._semantic_search(query, top_k * 2)
        keyword_results = self._keyword_search(query, top_k * 2)
        
        # Combine scores
        combined: Dict[str, float] = {}
        result_map: Dict[str, SearchResult] = {}
        
        for result in semantic_results:
            doc_id = result.document.id
            combined[doc_id] = self._config.semantic_weight * result.score
            result_map[doc_id] = result
        
        for result in keyword_results:
            doc_id = result.document.id
            current = combined.get(doc_id, 0)
            combined[doc_id] = current + self._config.keyword_weight * result.score
            
            if doc_id not in result_map:
                result_map[doc_id] = result
            else:
                # Merge highlights
                result_map[doc_id].highlights.extend(result.highlights)
        
        # Sort and return
        sorted_docs = sorted(combined.items(), key=lambda x: x[1], reverse=True)
        
        results = []
        for doc_id, score in sorted_docs[:top_k]:
            result = result_map[doc_id]
            result.score = score
            result.match_type = "hybrid"
            results.append(result)
        
        return results
    
    def _cosine_similarity(self, a: List[float], b: List[float]) -> float:
        """Compute cosine similarity."""
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x**2 for x in a) ** 0.5
        norm_b = sum(x**2 for x in b) ** 0.5
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot / (norm_a * norm_b)
    
    def answer(
        self,
        question: str,
        model: Any,
        top_k: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Answer a question using retrieved context.
        
        Args:
            question: Question to answer
            model: Language model
            top_k: Number of context documents
            
        Returns:
            Answer with sources
        """
        # Retrieve relevant documents
        results = self.query(question, top_k=top_k or 3)
        
        if not results:
            return {
                "answer": "I couldn't find relevant information to answer this question.",
                "sources": []
            }
        
        # Build context
        context_parts = []
        sources = []
        
        for result in results:
            context_parts.append(f"[{result.document.source}]\n{result.document.content}")
            sources.append({
                "source": result.document.source,
                "score": result.score
            })
        
        context = "\n\n".join(context_parts)
        
        # Generate answer
        prompt = f"""Answer the question based on the following context. If the context doesn't contain relevant information, say so.

Context:
{context}

Question: {question}

Answer:"""
        
        if hasattr(model, "generate"):
            answer = model.generate(prompt)
        else:
            answer = str(model(prompt))
        
        return {
            "answer": answer,
            "sources": sources,
            "context_used": len(results)
        }
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID."""
        return self._documents.get(doc_id)
    
    def list_sources(self) -> List[str]:
        """List all document sources."""
        sources = set()
        for doc in self._documents.values():
            sources.add(doc.source)
        return sorted(sources)
    
    def delete_source(self, source: str) -> int:
        """Delete all documents from a source."""
        to_delete = [
            doc_id for doc_id, doc in self._documents.items()
            if doc.source == source
        ]
        
        for doc_id in to_delete:
            del self._documents[doc_id]
        
        # Rebuild keyword index
        self._keyword_index = {}
        for doc in self._documents.values():
            self._index_keywords(doc)
        
        return len(to_delete)
    
    def _save(self):
        """Save to persistence directory."""
        if not self._persist_dir:
            return
        
        self._persist_dir.mkdir(parents=True, exist_ok=True)
        
        # Save documents
        doc_data = {
            doc_id: doc.to_dict()
            for doc_id, doc in self._documents.items()
        }
        
        with open(self._persist_dir / "documents.json", "w") as f:
            json.dump(doc_data, f)
    
    def _load(self):
        """Load from persistence directory."""
        doc_path = self._persist_dir / "documents.json"
        
        if doc_path.exists():
            with open(doc_path) as f:
                doc_data = json.load(f)
            
            for doc_id, data in doc_data.items():
                doc = Document(
                    id=data["id"],
                    content=data["content"],
                    source=data["source"],
                    metadata=data.get("metadata", {}),
                    chunk_index=data.get("chunk_index", 0),
                    total_chunks=data.get("total_chunks", 1)
                )
                self._documents[doc_id] = doc
                self._index_keywords(doc)
            
            logger.info(f"Loaded {len(self._documents)} documents")
    
    def stats(self) -> Dict[str, Any]:
        """Get system statistics."""
        return {
            "total_documents": len(self._documents),
            "total_sources": len(self.list_sources()),
            "total_keywords": len(self._keyword_index),
            "config": {
                "chunk_size": self._config.chunk_size,
                "chunk_overlap": self._config.chunk_overlap,
                "top_k": self._config.top_k
            }
        }


# Convenience function
def create_rag(persist_dir: Optional[str] = None) -> DocumentRAG:
    """Create a new RAG system."""
    return DocumentRAG(persist_dir=persist_dir)
