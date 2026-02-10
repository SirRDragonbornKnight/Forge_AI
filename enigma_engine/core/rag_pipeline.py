"""
================================================================================
RAG Pipeline - Retrieval-Augmented Generation
================================================================================

Built-in RAG for grounding LLM responses in your documents.
No need for LangChain or LlamaIndex - it's all included.

FEATURES:
    - Document loading (PDF, TXT, MD, DOCX)
    - Chunking with overlap
    - Vector embedding and indexing
    - Semantic search
    - Answer generation with sources

📍 FILE: enigma_engine/core/rag_pipeline.py
🏷️ TYPE: RAG System

USAGE:
    from enigma_engine.core.rag_pipeline import RAGPipeline
    
    # Create pipeline
    rag = RAGPipeline()
    
    # Index documents
    rag.add_document("docs/manual.pdf")
    rag.add_documents(["doc1.txt", "doc2.md"])
    
    # Query with RAG
    answer = rag.query("How do I configure the system?")
    print(answer.text)
    print(answer.sources)
"""

import hashlib
import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional, Union

logger = logging.getLogger(__name__)

# Optional imports
try:
    HAVE_TORCH = True
except ImportError:
    HAVE_TORCH = False

try:
    HAVE_NUMPY = True
except ImportError:
    HAVE_NUMPY = False


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class Document:
    """A document in the RAG system."""
    id: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    
    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")


@dataclass
class Chunk:
    """A chunk of a document."""
    id: str
    document_id: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    embedding: Optional[list[float]] = None
    
    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")


@dataclass
class SearchResult:
    """A search result with relevance score."""
    chunk: Chunk
    score: float


@dataclass
class RAGResponse:
    """Response from RAG query."""
    text: str
    sources: list[Chunk]
    context: str
    query: str


# =============================================================================
# Document Loader
# =============================================================================

class DocumentLoader:
    """Load documents from various formats."""
    
    @staticmethod
    def load(path: Union[str, Path]) -> str:
        """
        Load document content from file.
        
        Supported formats:
        - .txt, .md, .py, .json, etc. (text files)
        - .pdf (requires PyPDF2 or pdfplumber)
        - .docx (requires python-docx)
        """
        path = Path(path)
        suffix = path.suffix.lower()
        
        if suffix in ['.txt', '.md', '.py', '.json', '.yaml', '.yml', '.csv', '.html']:
            return DocumentLoader._load_text(path)
        elif suffix == '.pdf':
            return DocumentLoader._load_pdf(path)
        elif suffix == '.docx':
            return DocumentLoader._load_docx(path)
        else:
            # Try as text
            return DocumentLoader._load_text(path)
    
    @staticmethod
    def _load_text(path: Path) -> str:
        """Load text file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(path, encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode {path}")
    
    @staticmethod
    def _load_pdf(path: Path) -> str:
        """Load PDF file."""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n\n".join(pages)
        except ImportError:
            pass
        
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            pass
        
        raise ImportError("PDF loading requires pdfplumber or PyPDF2. "
                         "Install with: pip install pdfplumber")
    
    @staticmethod
    def _load_docx(path: Path) -> str:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("DOCX loading requires python-docx. "
                            "Install with: pip install python-docx")


# =============================================================================
# Text Chunker
# =============================================================================

class TextChunker:
    """Split text into overlapping chunks."""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separators: Optional[list[str]] = None
    ):
        """
        Initialize chunker.
        
        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            separators: Preferred split points (in order of preference)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]
    
    def chunk(self, text: str) -> list[str]:
        """Split text into chunks."""
        chunks = []
        
        # Recursive splitting
        def split_recursive(text: str, separators: list[str]) -> list[str]:
            if len(text) <= self.chunk_size:
                return [text] if text.strip() else []
            
            if not separators:
                # Force split at chunk_size
                return [
                    text[i:i + self.chunk_size]
                    for i in range(0, len(text), self.chunk_size - self.chunk_overlap)
                ]
            
            separator = separators[0]
            parts = text.split(separator)
            
            result = []
            current_chunk = ""
            
            for part in parts:
                test_chunk = current_chunk + separator + part if current_chunk else part
                
                if len(test_chunk) <= self.chunk_size:
                    current_chunk = test_chunk
                else:
                    if current_chunk:
                        result.append(current_chunk)
                    
                    if len(part) > self.chunk_size:
                        # Recursively split large parts
                        result.extend(split_recursive(part, separators[1:]))
                        current_chunk = ""
                    else:
                        current_chunk = part
            
            if current_chunk:
                result.append(current_chunk)
            
            return result
        
        raw_chunks = split_recursive(text, self.separators)
        
        # Add overlap
        for i, chunk in enumerate(raw_chunks):
            if i > 0 and self.chunk_overlap > 0:
                # Add end of previous chunk
                prev_chunk = raw_chunks[i - 1]
                overlap_text = prev_chunk[-self.chunk_overlap:]
                chunk = overlap_text + chunk
            chunks.append(chunk.strip())
        
        return [c for c in chunks if c]


# =============================================================================
# Simple Embedder (No Dependencies)
# =============================================================================

class SimpleEmbedder:
    """
    Simple embedding generator that works without external dependencies.
    
    Uses TF-IDF-like approach for reasonable semantic similarity.
    For better quality, use SentenceTransformerEmbedder.
    """
    
    def __init__(self, dim: int = 384):
        self.dim = dim
        self.vocab: dict[str, int] = {}
        self._idf: dict[str, float] = {}
    
    def _tokenize(self, text: str) -> list[str]:
        """Simple tokenization."""
        import re
        text = text.lower()
        words = re.findall(r'\b\w+\b', text)
        return words
    
    def _get_word_hash(self, word: str) -> int:
        """Get consistent hash for word."""
        return int(hashlib.md5(word.encode()).hexdigest()[:8], 16)
    
    def embed(self, text: str) -> list[float]:
        """Generate embedding for text."""
        tokens = self._tokenize(text)
        
        if not tokens:
            return [0.0] * self.dim
        
        # Create sparse embedding using hashing trick
        embedding = [0.0] * self.dim
        
        for token in tokens:
            idx = self._get_word_hash(token) % self.dim
            embedding[idx] += 1.0
        
        # Normalize
        norm = sum(x * x for x in embedding) ** 0.5
        if norm > 0:
            embedding = [x / norm for x in embedding]
        
        return embedding
    
    def embed_batch(self, texts: list[str]) -> list[list[float]]:
        """Embed multiple texts."""
        return [self.embed(text) for text in texts]


class SentenceTransformerEmbedder:
    """Embedder using sentence-transformers (high quality)."""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer(model_name)
            self.dim = self.model.get_sentence_embedding_dimension()
        except ImportError:
            raise ImportError(
                "SentenceTransformerEmbedder requires sentence-transformers. "
                "Install with: pip install sentence-transformers"
            )
    
    def embed(self, text: str) -> list[float]:
        """Generate embedding."""
        return self.model.encode(text, convert_to_numpy=True).tolist()
    
    def embed_batch(self, texts: list[str]) -> list[list[float]]:
        """Embed multiple texts."""
        embeddings = self.model.encode(texts, convert_to_numpy=True)
        return embeddings.tolist()


# =============================================================================
# Vector Store
# =============================================================================

class SimpleVectorStore:
    """
    Simple in-memory vector store with cosine similarity search.
    
    For production, consider using FAISS or a vector database.
    """
    
    def __init__(self):
        self.chunks: dict[str, Chunk] = {}
        self.embeddings: dict[str, list[float]] = {}
    
    def add(self, chunk: Chunk, embedding: list[float]):
        """Add a chunk with its embedding."""
        self.chunks[chunk.id] = chunk
        self.embeddings[chunk.id] = embedding
    
    def search(self, query_embedding: list[float], top_k: int = 5) -> list[SearchResult]:
        """Search for similar chunks."""
        if not self.chunks:
            return []
        
        results = []
        
        for chunk_id, embedding in self.embeddings.items():
            score = self._cosine_similarity(query_embedding, embedding)
            results.append(SearchResult(
                chunk=self.chunks[chunk_id],
                score=score
            ))
        
        # Sort by score (descending)
        results.sort(key=lambda x: x.score, reverse=True)
        
        return results[:top_k]
    
    def _cosine_similarity(self, a: list[float], b: list[float]) -> float:
        """Compute cosine similarity."""
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x * x for x in a) ** 0.5
        norm_b = sum(x * x for x in b) ** 0.5
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot / (norm_a * norm_b)
    
    def save(self, path: Union[str, Path]):
        """Save to file."""
        path = Path(path)
        
        data = {
            "chunks": {
                cid: {
                    "id": c.id,
                    "document_id": c.document_id,
                    "content": c.content,
                    "metadata": c.metadata
                }
                for cid, c in self.chunks.items()
            },
            "embeddings": self.embeddings
        }
        
        with open(path, 'w') as f:
            json.dump(data, f)
    
    def load(self, path: Union[str, Path]):
        """Load from file."""
        path = Path(path)
        
        with open(path) as f:
            data = json.load(f)
        
        self.chunks = {
            cid: Chunk(
                id=c["id"],
                document_id=c["document_id"],
                content=c["content"],
                metadata=c.get("metadata", {})
            )
            for cid, c in data["chunks"].items()
        }
        self.embeddings = data["embeddings"]
    
    def __len__(self):
        return len(self.chunks)


# =============================================================================
# RAG Pipeline
# =============================================================================

class RAGPipeline:
    """
    Complete RAG pipeline.
    
    Combines document loading, chunking, embedding, and retrieval
    into a simple interface.
    """
    
    def __init__(
        self,
        embedder: Optional[Any] = None,
        vector_store: Optional[SimpleVectorStore] = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        top_k: int = 5
    ):
        """
        Initialize RAG pipeline.
        
        Args:
            embedder: Embedding model (uses SimpleEmbedder if None)
            vector_store: Vector store (creates new if None)
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            top_k: Number of chunks to retrieve
        """
        # Try to use sentence-transformers, fall back to simple
        if embedder is None:
            try:
                self.embedder = SentenceTransformerEmbedder()
                logger.info("Using SentenceTransformer embedder")
            except ImportError:
                self.embedder = SimpleEmbedder()
                logger.info("Using simple embedder (install sentence-transformers for better quality)")
        else:
            self.embedder = embedder
        
        self.vector_store = vector_store or SimpleVectorStore()
        self.chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.top_k = top_k
        
        # Track documents
        self.documents: dict[str, Document] = {}
        
        # LLM for answer generation (lazy loaded)
        self._llm = None
    
    @property
    def llm(self):
        """Lazy-load LLM for answer generation."""
        if self._llm is None:
            try:
                from .inference import EnigmaEngine
                self._llm = EnigmaEngine()
            except Exception as e:
                logger.warning(f"Could not load LLM: {e}")
        return self._llm
    
    def add_document(
        self,
        path: Union[str, Path],
        metadata: Optional[dict[str, Any]] = None
    ) -> Document:
        """
        Add a document to the RAG system.
        
        Args:
            path: Path to document file
            metadata: Additional metadata
            
        Returns:
            The added Document
        """
        path = Path(path)
        
        # Load content
        content = DocumentLoader.load(path)
        
        # Create document
        doc_id = hashlib.md5(str(path).encode()).hexdigest()[:12]
        doc = Document(
            id=doc_id,
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                **(metadata or {})
            }
        )
        
        self.documents[doc_id] = doc
        
        # Chunk and embed
        chunks = self.chunker.chunk(content)
        
        for i, chunk_text in enumerate(chunks):
            chunk = Chunk(
                id=f"{doc_id}_{i}",
                document_id=doc_id,
                content=chunk_text,
                metadata={
                    "source": str(path),
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            
            # Generate embedding
            embedding = self.embedder.embed(chunk_text)
            
            # Add to vector store
            self.vector_store.add(chunk, embedding)
        
        logger.info(f"Added document: {path.name} ({len(chunks)} chunks)")
        return doc
    
    def add_documents(
        self,
        paths: list[Union[str, Path]],
        show_progress: bool = True
    ) -> list[Document]:
        """Add multiple documents."""
        documents = []
        
        for i, path in enumerate(paths):
            if show_progress:
                logger.info(f"Processing {i+1}/{len(paths)}: {path}")
            
            try:
                doc = self.add_document(path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Failed to add {path}: {e}")
        
        return documents
    
    def add_text(
        self,
        text: str,
        source: str = "manual",
        metadata: Optional[dict[str, Any]] = None
    ) -> Document:
        """
        Add raw text to the RAG system.
        
        Args:
            text: Text content
            source: Source identifier
            metadata: Additional metadata
            
        Returns:
            The added Document
        """
        doc_id = hashlib.md5(text[:100].encode()).hexdigest()[:12]
        doc = Document(
            id=doc_id,
            content=text,
            metadata={
                "source": source,
                **(metadata or {})
            }
        )
        
        self.documents[doc_id] = doc
        
        # Chunk and embed
        chunks = self.chunker.chunk(text)
        
        for i, chunk_text in enumerate(chunks):
            chunk = Chunk(
                id=f"{doc_id}_{i}",
                document_id=doc_id,
                content=chunk_text,
                metadata={
                    "source": source,
                    "chunk_index": i
                }
            )
            
            embedding = self.embedder.embed(chunk_text)
            self.vector_store.add(chunk, embedding)
        
        logger.info(f"Added text: {source} ({len(chunks)} chunks)")
        return doc
    
    def retrieve(self, query: str, top_k: Optional[int] = None) -> list[SearchResult]:
        """
        Retrieve relevant chunks for a query.
        
        Args:
            query: Search query
            top_k: Number of results (uses default if None)
            
        Returns:
            List of SearchResult
        """
        top_k = top_k or self.top_k
        
        # Embed query
        query_embedding = self.embedder.embed(query)
        
        # Search
        results = self.vector_store.search(query_embedding, top_k=top_k)
        
        return results
    
    def query(
        self,
        question: str,
        top_k: Optional[int] = None,
        include_sources: bool = True,
        max_tokens: int = 256
    ) -> RAGResponse:
        """
        Query the RAG system with a question.
        
        Args:
            question: The question to answer
            top_k: Number of context chunks
            include_sources: Include source references
            max_tokens: Max tokens for answer generation
            
        Returns:
            RAGResponse with answer and sources
        """
        # Retrieve relevant chunks
        results = self.retrieve(question, top_k=top_k)
        
        if not results:
            return RAGResponse(
                text="I don't have any relevant information to answer this question.",
                sources=[],
                context="",
                query=question
            )
        
        # Build context
        context_parts = []
        for i, result in enumerate(results):
            context_parts.append(f"[{i+1}] {result.chunk.content}")
        
        context = "\n\n".join(context_parts)
        
        # Generate answer with LLM
        if self.llm:
            prompt = f"""Answer the question based on the following context. If the context doesn't contain relevant information, say so.

Context:
{context}

Question: {question}

Answer:"""
            
            try:
                answer = self.llm.generate(
                    prompt,
                    max_gen=max_tokens,
                    temperature=0.7
                )
            except Exception as e:
                logger.error(f"LLM generation failed: {e}")
                answer = f"Based on the retrieved context, here are the relevant excerpts:\n\n{context}"
        else:
            # No LLM - return context directly
            answer = f"Here are the most relevant excerpts:\n\n{context}"
        
        return RAGResponse(
            text=answer,
            sources=[r.chunk for r in results],
            context=context,
            query=question
        )
    
    def save(self, path: Union[str, Path]):
        """Save the RAG index to disk."""
        path = Path(path)
        path.mkdir(parents=True, exist_ok=True)
        
        # Save vector store
        self.vector_store.save(path / "vectors.json")
        
        # Save documents
        docs_data = {
            doc_id: {
                "id": doc.id,
                "content": doc.content,
                "metadata": doc.metadata
            }
            for doc_id, doc in self.documents.items()
        }
        
        with open(path / "documents.json", 'w') as f:
            json.dump(docs_data, f)
        
        logger.info(f"RAG index saved to {path}")
    
    def load(self, path: Union[str, Path]):
        """Load a RAG index from disk."""
        path = Path(path)
        
        # Load vector store
        self.vector_store.load(path / "vectors.json")
        
        # Load documents
        with open(path / "documents.json") as f:
            docs_data = json.load(f)
        
        self.documents = {
            doc_id: Document(
                id=d["id"],
                content=d["content"],
                metadata=d.get("metadata", {})
            )
            for doc_id, d in docs_data.items()
        }
        
        logger.info(f"RAG index loaded from {path} ({len(self.documents)} documents)")
    
    def stats(self) -> dict[str, Any]:
        """Get RAG system statistics."""
        return {
            "documents": len(self.documents),
            "chunks": len(self.vector_store),
            "embedder": type(self.embedder).__name__,
        }


# =============================================================================
# Factory Functions
# =============================================================================

def create_rag_pipeline(
    documents: Optional[list[Union[str, Path]]] = None,
    use_gpu: bool = True,
    **kwargs
) -> RAGPipeline:
    """
    Create a RAG pipeline, optionally indexing documents.
    
    Args:
        documents: List of document paths to index
        use_gpu: Use GPU for embeddings if available
        **kwargs: Additional pipeline arguments
        
    Returns:
        Configured RAGPipeline
    """
    pipeline = RAGPipeline(**kwargs)
    
    if documents:
        pipeline.add_documents(documents)
    
    return pipeline


# Example usage
if __name__ == "__main__":
    print("RAG Pipeline Demo")
    print("=" * 50)
    
    # Create pipeline
    rag = RAGPipeline(chunk_size=200)
    
    # Add some sample text
    rag.add_text(
        """
        Enigma AI Engine is an AI framework for building and running language models.
        It supports training from scratch, fine-tuning, and inference.
        
        Key features include:
        - Multiple model sizes from nano to XXL
        - GPU and CPU support
        - Built-in tokenizer training
        - API server for deployment
        
        To train a model, use: python run.py --train
        To run inference: python run.py --run
        """,
        source="docs"
    )
    
    print(f"RAG stats: {rag.stats()}")
    
    # Query
    results = rag.retrieve("How do I train a model?", top_k=2)
    
    print(f"\nQuery: 'How do I train a model?'")
    print(f"Found {len(results)} relevant chunks:")
    for i, result in enumerate(results):
        print(f"\n[{i+1}] Score: {result.score:.3f}")
        print(f"    {result.chunk.content[:100]}...")
