"""
Document Ingestion Pipeline for Enigma AI Engine

Process and ingest various document formats.

Features:
- PDF extraction
- Word document parsing
- Web scraping
- Text chunking
- Metadata extraction
- Content cleaning

Usage:
    from enigma_engine.tools.document_ingestion import DocumentIngester, ChunkingConfig
    
    ingester = DocumentIngester()
    
    # Ingest a document
    chunks = ingester.ingest("document.pdf")
    
    # Ingest from URL
    chunks = ingester.ingest_url("https://example.com/page")
"""

import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Callable, Dict, Iterator, List, Optional, Tuple
from urllib.parse import urlparse

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MARKDOWN = "markdown"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    EPUB = "epub"
    RTF = "rtf"


@dataclass
class ChunkingConfig:
    """Configuration for text chunking."""
    chunk_size: int = 1000  # Characters per chunk
    chunk_overlap: int = 200  # Overlap between chunks
    min_chunk_size: int = 100  # Minimum chunk size
    split_on_sentences: bool = True  # Prefer sentence boundaries
    split_on_paragraphs: bool = True  # Prefer paragraph boundaries
    preserve_markdown: bool = True  # Keep markdown formatting


@dataclass
class DocumentMetadata:
    """Metadata for a document."""
    source: str
    title: str = ""
    author: str = ""
    created_date: str = ""
    modified_date: str = ""
    num_pages: int = 0
    file_size_bytes: int = 0
    content_hash: str = ""
    document_type: DocumentType = DocumentType.TXT
    language: str = "en"
    tags: List[str] = field(default_factory=list)
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentChunk:
    """A chunk of a document."""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: DocumentMetadata = None
    embeddings: Optional[List[float]] = None


class TextExtractor:
    """Extract text from various document formats."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Extract text from file.
        
        Returns:
            Tuple of (text, metadata)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        metadata = DocumentMetadata(
            source=str(file_path),
            file_size_bytes=file_path.stat().st_size
        )
        
        if suffix == '.pdf':
            text, metadata = self._extract_pdf(file_path, metadata)
        elif suffix == '.docx':
            text, metadata = self._extract_docx(file_path, metadata)
        elif suffix in ['.txt', '.text']:
            text, metadata = self._extract_text(file_path, metadata)
        elif suffix in ['.html', '.htm']:
            text, metadata = self._extract_html(file_path, metadata)
        elif suffix in ['.md', '.markdown']:
            text, metadata = self._extract_markdown(file_path, metadata)
        elif suffix == '.csv':
            text, metadata = self._extract_csv(file_path, metadata)
        elif suffix == '.json':
            text, metadata = self._extract_json(file_path, metadata)
        else:
            # Try as plain text
            text, metadata = self._extract_text(file_path, metadata)
        
        metadata.content_hash = hashlib.sha256(text.encode()).hexdigest()
        
        return text, metadata
    
    def _extract_pdf(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text from PDF."""
        metadata.document_type = DocumentType.PDF
        
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata.num_pages = len(reader.pages)
                
                # Extract metadata
                if reader.metadata:
                    metadata.title = reader.metadata.get('/Title', '')
                    metadata.author = reader.metadata.get('/Author', '')
                
                for page in reader.pages:
                    text_parts.append(page.extract_text())
            
            return '\n\n'.join(text_parts), metadata
            
        except ImportError:
            logger.warning("PyPDF2 not installed, trying pdfminer")
            
            try:
                from pdfminer.high_level import extract_text
                text = extract_text(file_path)
                return text, metadata
            except ImportError:
                raise ImportError("No PDF library available. Install PyPDF2 or pdfminer.six")
    
    def _extract_docx(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text from Word document."""
        metadata.document_type = DocumentType.DOCX
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract metadata
            if doc.core_properties:
                metadata.title = doc.core_properties.title or ''
                metadata.author = doc.core_properties.author or ''
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text_parts.append(' | '.join(cells))
            
            return '\n\n'.join(text_parts), metadata
            
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    def _extract_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract from plain text file."""
        metadata.document_type = DocumentType.TXT
        metadata.title = file_path.stem
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_path.read_text(encoding=encoding)
                return text, metadata
            except UnicodeDecodeError:
                continue
        
        # Fallback: read as bytes and decode with errors='replace'
        text = file_path.read_bytes().decode('utf-8', errors='replace')
        return text, metadata
    
    def _extract_html(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text from HTML."""
        metadata.document_type = DocumentType.HTML
        
        html = file_path.read_text()
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Get title
            if soup.title:
                metadata.title = soup.title.string or ''
            
            # Remove script and style elements
            for element in soup(['script', 'style', 'nav', 'footer', 'header']):
                element.decompose()
            
            text = soup.get_text(separator='\n')
            
        except ImportError:
            # Fallback: regex-based extraction
            text = re.sub(r'<script.*?</script>', '', html, flags=re.DOTALL)
            text = re.sub(r'<style.*?</style>', '', text, flags=re.DOTALL)
            text = re.sub(r'<[^>]+>', '', text)
            text = re.sub(r'&\w+;', ' ', text)
        
        # Clean up whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        return text.strip(), metadata
    
    def _extract_markdown(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract from Markdown."""
        metadata.document_type = DocumentType.MARKDOWN
        
        text = file_path.read_text()
        
        # Extract title from first heading
        title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if title_match:
            metadata.title = title_match.group(1)
        
        return text, metadata
    
    def _extract_csv(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract from CSV."""
        metadata.document_type = DocumentType.CSV
        metadata.title = file_path.stem
        
        import csv
        
        text_parts = []
        with open(file_path, 'r', newline='') as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(' | '.join(row))
        
        return '\n'.join(text_parts), metadata
    
    def _extract_json(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, DocumentMetadata]:
        """Extract from JSON."""
        import json
        
        metadata.document_type = DocumentType.JSON
        metadata.title = file_path.stem
        
        data = json.loads(file_path.read_text())
        
        # Convert to readable text
        text = json.dumps(data, indent=2)
        return text, metadata


class WebExtractor:
    """Extract text from web pages."""
    
    def __init__(
        self,
        user_agent: str = "EnigmaBot/1.0",
        timeout: int = 30
    ):
        self.user_agent = user_agent
        self.timeout = timeout
    
    def extract(self, url: str) -> Tuple[str, DocumentMetadata]:
        """
        Extract text from URL.
        
        Returns:
            Tuple of (text, metadata)
        """
        metadata = DocumentMetadata(
            source=url,
            document_type=DocumentType.HTML
        )
        
        try:
            import requests
            
            headers = {'User-Agent': self.user_agent}
            response = requests.get(url, headers=headers, timeout=self.timeout)
            response.raise_for_status()
            
            html = response.text
            
        except ImportError:
            # Fallback to urllib
            import urllib.request
            
            req = urllib.request.Request(url, headers={'User-Agent': self.user_agent})
            with urllib.request.urlopen(req, timeout=self.timeout) as response:
                html = response.read().decode('utf-8')
        
        # Parse HTML
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Get title
            if soup.title:
                metadata.title = soup.title.string or ''
            
            # Get main content
            main = soup.find('main') or soup.find('article') or soup.body
            
            if main:
                # Remove navigation, footer, etc.
                for element in main.find_all(['nav', 'footer', 'aside', 'script', 'style']):
                    element.decompose()
                
                text = main.get_text(separator='\n')
            else:
                text = soup.get_text(separator='\n')
            
        except ImportError:
            # Fallback
            text = re.sub(r'<[^>]+>', '', html)
        
        # Clean
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        metadata.content_hash = hashlib.sha256(text.encode()).hexdigest()
        
        return text, metadata


class TextChunker:
    """Split text into chunks."""
    
    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or ChunkingConfig()
    
    def chunk(
        self,
        text: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> List[DocumentChunk]:
        """
        Split text into chunks.
        
        Returns:
            List of DocumentChunk
        """
        if not text:
            return []
        
        # Split into paragraphs first
        if self.config.split_on_paragraphs:
            paragraphs = re.split(r'\n\s*\n', text)
        else:
            paragraphs = [text]
        
        chunks = []
        current_chunk = ""
        current_start = 0
        char_pos = 0
        chunk_index = 0
        
        for para in paragraphs:
            # Would adding this paragraph exceed chunk size?
            if current_chunk and len(current_chunk) + len(para) + 2 > self.config.chunk_size:
                # Save current chunk
                if len(current_chunk) >= self.config.min_chunk_size:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=char_pos,
                        metadata=metadata
                    ))
                    chunk_index += 1
                
                # Start new chunk with overlap
                if self.config.chunk_overlap > 0:
                    overlap_text = current_chunk[-self.config.chunk_overlap:]
                    current_chunk = overlap_text + "\n\n" + para
                    current_start = char_pos - len(overlap_text)
                else:
                    current_chunk = para
                    current_start = char_pos
            else:
                if current_chunk:
                    current_chunk += "\n\n"
                current_chunk += para
            
            char_pos += len(para) + 2  # +2 for paragraph separator
        
        # Don't forget last chunk
        if current_chunk and len(current_chunk) >= self.config.min_chunk_size:
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=char_pos,
                metadata=metadata
            ))
        
        return chunks
    
    def chunk_by_sentences(
        self,
        text: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> List[DocumentChunk]:
        """Chunk text by sentences."""
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        current_start = 0
        char_pos = 0
        chunk_index = 0
        
        for sentence in sentences:
            if current_chunk and len(current_chunk) + len(sentence) + 1 > self.config.chunk_size:
                if len(current_chunk) >= self.config.min_chunk_size:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=char_pos,
                        metadata=metadata
                    ))
                    chunk_index += 1
                
                current_chunk = sentence
                current_start = char_pos
            else:
                if current_chunk:
                    current_chunk += " "
                current_chunk += sentence
            
            char_pos += len(sentence) + 1
        
        if current_chunk and len(current_chunk) >= self.config.min_chunk_size:
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=char_pos,
                metadata=metadata
            ))
        
        return chunks


class DocumentIngester:
    """Complete document ingestion pipeline."""
    
    def __init__(
        self,
        chunking_config: Optional[ChunkingConfig] = None,
        embedding_fn: Optional[Callable[[str], List[float]]] = None
    ):
        """
        Initialize document ingester.
        
        Args:
            chunking_config: Configuration for chunking
            embedding_fn: Function to generate embeddings
        """
        self.text_extractor = TextExtractor()
        self.web_extractor = WebExtractor()
        self.chunker = TextChunker(chunking_config)
        self.embedding_fn = embedding_fn
        
        logger.info("DocumentIngester initialized")
    
    def ingest(
        self,
        source: str,
        generate_embeddings: bool = False
    ) -> List[DocumentChunk]:
        """
        Ingest a document from file path.
        
        Args:
            source: File path
            generate_embeddings: Whether to generate embeddings
            
        Returns:
            List of document chunks
        """
        # Extract text
        text, metadata = self.text_extractor.extract(source)
        
        # Chunk
        chunks = self.chunker.chunk(text, metadata)
        
        # Generate embeddings
        if generate_embeddings and self.embedding_fn:
            for chunk in chunks:
                chunk.embeddings = self.embedding_fn(chunk.content)
        
        logger.info(f"Ingested {source}: {len(chunks)} chunks")
        
        return chunks
    
    def ingest_url(
        self,
        url: str,
        generate_embeddings: bool = False
    ) -> List[DocumentChunk]:
        """
        Ingest a document from URL.
        
        Returns:
            List of document chunks
        """
        text, metadata = self.web_extractor.extract(url)
        chunks = self.chunker.chunk(text, metadata)
        
        if generate_embeddings and self.embedding_fn:
            for chunk in chunks:
                chunk.embeddings = self.embedding_fn(chunk.content)
        
        logger.info(f"Ingested {url}: {len(chunks)} chunks")
        
        return chunks
    
    def ingest_directory(
        self,
        directory: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True
    ) -> List[DocumentChunk]:
        """
        Ingest all documents in a directory.
        
        Args:
            directory: Directory path
            extensions: File extensions to include
            recursive: Search recursively
            
        Returns:
            Combined list of chunks
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        extensions = extensions or ['.txt', '.md', '.pdf', '.docx', '.html']
        
        all_chunks = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    chunks = self.ingest(str(file_path))
                    all_chunks.extend(chunks)
                except Exception as e:
                    logger.error(f"Failed to ingest {file_path}: {e}")
        
        logger.info(f"Ingested {len(all_chunks)} chunks from {directory}")
        
        return all_chunks
    
    def ingest_batch(
        self,
        sources: List[str],
        generate_embeddings: bool = False
    ) -> Dict[str, List[DocumentChunk]]:
        """
        Ingest multiple documents.
        
        Returns:
            Dict of source -> chunks
        """
        results = {}
        
        for source in sources:
            try:
                if source.startswith(('http://', 'https://')):
                    chunks = self.ingest_url(source, generate_embeddings)
                else:
                    chunks = self.ingest(source, generate_embeddings)
                results[source] = chunks
            except Exception as e:
                logger.error(f"Failed to ingest {source}: {e}")
                results[source] = []
        
        return results
